"""
LOI (Letter of Intent) Service
Generates downloadable LOI documents (PDF + DOCX) from PO data.
No DB persistence - purely a form-time document generator.

Usage:
    from app.services.loi_service import LOIService
    filled = LOIService.fill_template("technology", po_data)
    pdf_bytes = LOIService.generate_pdf(filled, edited_articles)
    docx_bytes = LOIService.generate_docx(filled, edited_articles)
"""
from __future__ import annotations
from io import BytesIO
from typing import Any


# ── LOI Article Templates ─────────────────────────────────────────────────────
# Placeholders:
#   {vendor_name}        - vendor's company name
#   {total_amount}       - PO total amount in INR (numeric)
#   {total_amount_words} - amount in words (e.g. "Nine Crore Forty Two Lakhs only")
#   {description}        - PO description / purpose
#   {delivery_address}   - delivery address from PO
#   {required_by}        - required by date
#   {penalty_clauses}    - from PO penalty_clauses field
#   {delivery_terms}     - from PO delivery_terms field
#   {warranty_terms}     - from PO warranty_terms field
#   {special_conditions} - from PO special_conditions field
#   {vendor_contact}     - vendor contact name (from PO payment_terms field for now)
#   {site_name}          - site name from PO
#   {po_number}          - PO number

TECHNOLOGY_ARTICLES = [
    {"number": "1", "title": "SCOPE OF WORK", "body": "M/s {vendor_name} will provide Complete Design, Engineering, Manufacturing and Supply components as per Annexure-I and tender specifications to meet the outlet parameter as per the Tender Requirements.\n\nScope of supply for this project:\n\n{description}"},
    {"number": "2", "title": "BASIS OF DESIGN AND TECHNICAL REQUIREMENT", "body": "The system shall be designed, including all related accessories, for the technical requirements as specified in the tender documents and as agreed upon between PEEIPL and {vendor_name}. {vendor_name} shall assume suitable concentration for components not given in the table and shall consider all relevant parameters while designing the system.\n\n{description}"},
    {"number": "3", "title": "PROCESS DESCRIPTION", "body": "The system has the following units for treatment as per project requirements:\n\n{description}\n\nAll process parameters and performance guarantees shall be as per the tender requirements and mutually agreed specifications."},
    {"number": "4", "title": "EXCLUSIONS", "body": "1. All civil works activities shall be in scope of PEEIPL.\n2. Any items not explicitly mentioned in the scope of supply.\n3. {special_conditions}"},
    {"number": "5", "title": "ENGINEERING", "body": "{vendor_name} shall submit the following engineering documents: All engineering support during Basic Engineering Package including Drawings, P&ID, Process description, GA, PLC, Operational Philosophy, Control Philosophy, I/O list, Electrical and Instrumentation details, Material submittal, Auto cad drawings, Civil input drawings consisting of general arrangement drawings showing inside dimensions of concrete tanks only, Electrical input drawings, method statement of installation and testing commissioning, Quality Assurance plan, Operation Maintenance manual (editable soft copy as well hard copy) etc. within 8-10 weeks after receipt of fully signed LOI."},
    {"number": "6", "title": "VALIDATION", "body": "{vendor_name} will review Contractor's drawings in context of the system for conformance to {vendor_name}'s requirements."},
    {"number": "7", "title": "PRICE", "body": "The total cost for the Complete Design, Engineering, Manufacturing, Supply, Supervision for Installation, Testing & Commissioning shall be INR {total_amount} ({total_amount_words}) including custom duty, port clearance charges, transit insurance and transportation to site.\n\nGST shall be paid extra as applicable."},
    {"number": "8", "title": "PAYMENT TERMS", "body": "1. 90% of PO value against LC opened on any nationalised Indian Bank with usance period of 75 days from the date of Bill of Exchange.\n2. 10% of PO value within 15 days of commissioning and PG test of system subject to submission of PBG of 10% order value valid for 7 years from the date of commissioning.\n3. LC will be opened and handed over along with manufacturing clearance.\n4. Charges with respect to opening of LC will be borne by PEEIPL and all other charges will be borne by {vendor_name}.\n\nPricing Notes: All prices quoted are in INR\n\n{vendor_name}'s Guarantee of title: {vendor_name} warrants and guarantees that title to all materials and equipment will pass to BUYER free and clear of all Liens."},
    {"number": "9", "title": "DELIVERY", "body": "Kick of Meeting: Within 7 days after receipt of fully signed LOI\n\nBasic Engineering Submittals: Within 10 days after Kick of meeting.\n\nDetailed Engineering Submittals: Within 6 weeks from receipt of client's approval on basic engineering documents\n\nClient's Review & approval: Within 30 Days upon receipt of engineering documents by {vendor_name}.\n\nManufacturing & Supply: {delivery_terms}\n\nDelivery Address: {delivery_address}\nRequired By: {required_by}"},
    {"number": "10", "title": "UNDERTAKING", "body": "It is also understood and agreed upon that {vendor_name} will supply the material and services as per the approved submittals.\n\n{vendor_name} will support Passavant for approval of design engineering documents with the client.\n\nThe Buyer is obligated to inspect the materials at site. Any third party fees towards material inspection by Buyer/Buyer's client at Supplier/Supplier's manufacturer's facility shall be borne by the Buyer, in case required."},
    {"number": "11", "title": "PACKING", "body": "Packing/Containerization shall be in accordance with International Standards."},
    {"number": "12A", "title": "MECHANICAL WARRANTY", "body": "Mechanical Warranty of supplied equipment shall be valid for 12 months from the date of commissioning and handing over or 18 months from the date of last shipment whichever is earlier.\n\n{warranty_terms}"},
    {"number": "12B", "title": "PERFORMANCE GUARANTEE", "body": "Performance guarantee for the system specified in the contract Article 2 for 7 years of operation and maintenance after commissioning of equipment at site."},
    {"number": "13", "title": "TECHNICAL SERVICES", "body": "{vendor_name} should depute experience service engineer/representative for Installation inspection and testing & commissioning as per site requirement. PEEIPL will give 2 weeks' notice to {vendor_name} for deputing their service engineer at site. Total duration of visits required at site shall be 45 mandays for each plant separately.\n\n* {vendor_name} also needs to confirm their experienced local representative available within 48 Hrs. in case of any urgency during warranty period.\n* {vendor_name} should send engineering representative for clarification and immediate approval from client on request from PEEIPL and client without any charges."},
    {"number": "14", "title": "PROCESS GUARANTEE", "body": "{vendor_name} shall guarantee the following outlet parameters at the outlet of the System during the 7 years O&M from testing and commissioning of plant as per provided parameters subject to plant Operation is done as per the O&M Manual.\n\n{vendor_name} shall perform as per their final offer and as per tender requirements.\n\nDamages during Test on Completion, Damages shall not exceed 2% of that plant's PEEIPL Contract Price."},
    {"number": "15", "title": "TRAINING", "body": "{vendor_name} is responsible to provide training to ultimate client and operators of the plant during startup and commissioning as per site requirement."},
    {"number": "16", "title": "RESPONSIBILITY", "body": "{vendor_name} must accept the commitment and responsibility for the purchased product to be in full compliance with {vendor_name}'s Standards.\n\n16.1 Materials and Equipment: All materials and equipment incorporated into the work shall be of good quality and new, except as otherwise provided in the Contract Documents.\n\n16.2 No sub suppliers/sub-contractors shall be engaged without the prior approval of PEEIPL unless it is as per approved vendor list. {vendor_name} shall be responsible to BUYER for the acts and omissions of its subcontractor's suppliers and other individuals or entities."},
    {"number": "17", "title": "SELLER'S DEFAULT", "body": "17.1 The Seller is considered in default in case any or all of the following events:-\n17.1.1 Fails to proceed with the purchase order with due diligence after being required to do so in writing.\n17.1.2 Fails to execute promptly his obligations in accordance with the Purchase Order after being requested in writing by the Buyer.\n17.1.3 Fails to remove defective materials or make good defective work after being directed in writing to do so.\n17.1.4 Commits an act of bankruptcy or goes into liquidation.\n17.1.5 Fails to successfully execute, deliver and complete the supply within stipulated time.\n17.1.6 The Seller causes unreasonable delay to the project due to lack of materials or any reason exclusive from force majeure.\n\n17.2.2 In addition, the buyer is entitled to impose Liquidated damage to the Seller at the rate of 0.5% (Zero Point Five Per Cent) of Ex work price on the contract value for each week of delay or any part thereof up to a maximum of 10% of Ex work price of contract value for time at large as a default of the Seller."},
    {"number": "18", "title": "BUYER'S DEFAULT", "body": "18.1 The buyer is considered in default in case any or all of following events-\n18.1.1 Fails to provide manufacturing clearance, after being required to do so as per agreed terms & conditions and mutually agreed project schedule in writing.\n18.1.2 Fails to make payment, Open LC in accordance with purchase order after being requested by seller after successful inspection prior to dispatch.\n18.1.3 Commits and act of Bankruptcy or goes into liquidation.\n18.1.4 Fails to issue dispatch clearance upto 04 months from date of inspection.\n\nNote: Buyer agrees to reimburse actual storage charges in India for storage period post that period buyer allows seller to divert the material."},
    {"number": "19", "title": "DELAY", "body": "19.1 In case of any of the following events:-\n19.1.1 The Seller delays delivery of products and any associated service as stipulated in the Contract.\n19.1.2 The Seller delivers goods partially thereby delay the buyer's schedule.\n19.1.3 The Seller delivers goods not in accordance with the Purchase Order.\n19.1.4 The Seller delivers unapproved goods to agreed port and delays replacement of the approved materials.\n\nIn case, the seller performance is delayed due to any delay on client's side, the seller shall be given due extension of time without any cost implication.\n\n19.2.1 The buyer is entitled to pursue the outstanding materials and deduct the amount incurred from any payment outstanding to the Seller.\n\n19.2.2 The buyer is entitled to recover all reasonable costs and damage incurred by the Buyer / Client as a result of the Seller's delay, vide outstanding payments due to the Seller, liquidation of the Seller's Performance Guarantee or any other means at the buyer's disposal."},
    {"number": "20", "title": "LIQUIDATED DAMAGES", "body": "If the Seller fails to successfully execute, deliver and complete the supply within stipulated time as mentioned in the contract, for the reasons solely attributable to Seller's default and subject to force majeure, then the seller shall be liable to pay buyer Liquidated damages, a sum equivalent to 0.5% (Zero Point Five Per Cent) of the Purchase order value excluding taxes and duties for each week of delay or any part thereof. However, total amount of liquidated damage for delay in completion of contract shall be subject to a maximum of 10% of purchase order price excluding taxes and duties.\n\n{penalty_clauses}\n\nArticle 18 will be applicable only after the LD period is over."},
    {"number": "21", "title": "SUSPENSION OF WORKS", "body": "In the event of noncompliance or breach of any of the terms and conditions of the contract or the material default of the works under this contract the non-defaulting party shall furnish notice to defaulting party and in the event the defaulting party fails to cure the breach within 21 days of the notice, the non-defaulting party shall be at liberty to revoke/suspend the contract.\n\nThe buyer shall be at liberty to suspend the contract for convenience by written notice to the seller, to suspend the work not more than 24 months."},
    {"number": "22", "title": "TERMINATION", "body": "Either party (Non Defaulter Party) may terminate the contract or this agreement for default to the party (Default party) commits a material breach of this agreement and fails to cure the breach (If curable) within 21 days from the date of notice from the non-defaulting party.\n\nUpon the terminations of this agreement by the buyer under this provision, (i) Buyer shall order the goods and services himself, or employ the Seller at the expense of the seller. (ii) Seller shall reimburse Buyer the difference between that portion of the agreement price allocable to the terminated scope and the actual amounts reasonably incurred by the Buyer to complete the scope.\n\nUpon the termination of this agreement by seller under the provision of proceeding Article 18.1.2 of this LOI (i) Buyer shall pay to seller 50% of Ex-works value of this contract."},
    {"number": "22A", "title": "ARBITRATION", "body": "In case any dispute relating to the terms and conditions of this Contract or the interpretation thereof arises between the parties, the same shall promptly and in good faith be negotiated with a view of its amicable resolution and settlement.\n\nIn the event no amicable resolution or settlement is reached within a period of 30 days from the day on which the dispute(s) or difference(s) arose, such dispute(s) or difference(s) shall be referred to and settled by the arbitration. Three arbitrators shall be appointed - one each by the Buyer and Seller and third arbitrator shall be jointly appointed by the two Arbitrators.\n\n- The place of arbitration shall be Delhi.\n- The decision and award resulting from such arbitration shall be final and binding on the Parties."},
    {"number": "23", "title": "MISCELLANEOUS", "body": "PEEIPL to provide all relevant specifications and sections to the seller. {vendor_name}'s representative will be available by phone or through video (Skype) conference for project kick off meeting and for any engineering interface with ultimate client.\n\n23.1 {vendor_name} will assign Project Manager for this project who will be responsible for the primary contact for Buyer throughout the order execution & system commissioning."},
    {"number": "24", "title": "GENERAL INDEMNITY", "body": "Seller shall indemnify and hold harmless Buyer from claims for physical damage to third party property or injury to persons, including death, to the extent caused by the negligence of Seller or its officers, agents, employees, and/or assigns while engaged in activities under this Agreement."},
    {"number": "25", "title": "FORCE MAJURE", "body": "Seller shall not be liable nor in breach or default of its obligations under this Agreement to the extent performance of such obligations is delayed or prevented, directly or indirectly, due to causes beyond the reasonable control of Seller, including, but not limited to: acts of God, natural disasters, any act (or omission) by any governmental authority."},
    {"number": "26", "title": "CONFIDENTIALITY, INTELLECTUAL PROPERTY", "body": "Both Parties agree to keep confidential the other Party's proprietary non-public information, if any, which may be acquired in connection with this Agreement. Seller retains all intellectual property rights including copyright which it has in all drawings and data or other deliverables supplied or developed under this Agreement.\n\nAny software Seller owns and provides pursuant to this Agreement shall remain Seller's property. Seller provides to Buyer a limited, non-exclusive and terminable royalty free project-specific license to such software for the use, operation or maintenance at Buyer's site."},
    {"number": "27", "title": "LIMITATION OF LIABILITIES", "body": "The aggregate liability of the seller with respect to all claims arising out of or in connection with performance or non-performance of this contract whether in contract, warranty, tort or otherwise shall not exceed the contract price provided that this limitation shall not apply in case of negligence, willful misconduct or liabilities arising out of indemnity provisions in this contract.\n\n1) In no event other than buyer's default, seller shall be liable for any loss of profit or revenues, loss of production, loss of use of equipment or services or any associated equipment, interruption of business, or for any special, consequential incidental, indirect, punitive or exemplary damages.\n\n2) Seller's liability shall end upon expiration of the applicable warranty period."},
    {"number": "28", "title": "JURISDICTION", "body": "All suits, legal proceedings and arbitration award under this Contract shall be filed, entertained and decided in the Court of Delhi and the Courts in Delhi shall have the exclusive jurisdiction over all such disputes/claims."},
]

# Service and Supply use same structure with different scope/process text
# (will be filled in when client provides the actual text)
# SERVICE_ARTICLES_UPDATED -- 42-clause structure for civil/construction service POs
SERVICE_ARTICLES = [
    {
        "number": "1",
        "title": "Scope of Work",
        "body": "The Scope of Work is described in the annexures attached. {vendor_name} is advised that the scope mentioned is not limited to these annexures only, but also bound and executed with all terms and conditions of tender & agreement between PEEIPL and {vendor_name}.\n\nThe various jobs are to be executed as per Standard Specifications and guidelines of the Client/PEEIPL, which are forming part of the Contract between {vendor_name} and PEEIPL.\n\nAny work not specifically mentioned in the scope of work, but necessary for the satisfactory execution and completion of the assigned jobs, is deemed to be included in the Scope of Work to be executed by {vendor_name} within the specified time and the Work/Scope/contract value.\n\n{description}"
    },
    {
        "number": "2",
        "title": "Site Visit",
        "body": "{vendor_name} has already visited the site and made himself well acquainted with the location, area and related logistics and equipments for completing the specified works.\n\nSite: {site_name}\nDelivery Address: {delivery_address}"
    },
    {
        "number": "3",
        "title": "Value of the Work/Scope/Contract",
        "body": "The value of this Contract shall be INR {total_amount} ({total_amount_words}).\n\nThe payments shall be made by PEEIPL based on the unit rates and quantities and as per measurement of actual work executed at the project site.\n\nAny deduction made by the client due to poor performance of work by {vendor_name}, the same amount will be recovered from {vendor_name}'s bills."
    },
    {
        "number": "4",
        "title": "Rates and Escalation",
        "body": "The unit rates agreed upon shall remain firm and unchanged throughout the total period of execution of works under this work including the extended period, if any. This being firm and final price, no escalation is payable in future on any account whatsoever it may be."
    },
    {
        "number": "5",
        "title": "Income Tax",
        "body": "Income Tax (TDS), TCS and Labour Cess shall be deducted at applicable rate in accordance with the relevant laws from all payments made by PEEIPL. TDS certificates for the same shall be provided by PEEIPL."
    },
    {
        "number": "6",
        "title": "Royalties",
        "body": "{vendor_name} shall pay required royalties and fees, wherever applicable and shall submit proof/receipt of all such payments with PEEIPL along with Running Account (RA) Bills/invoices. {vendor_name} shall also procure, as required, all appropriate proprietary rights, licenses, agreements and permissions for materials, methods, processes, intellectual property incorporated into the works."
    },
    {
        "number": "7",
        "title": "Terms of Payment",
        "body": "7.1 Payment process: The Works shall be measured and remunerated according to the Main Contract. The BOQ rates are all inclusive for the execution of the Works by {vendor_name}.\n\n7.2 Payment of monthly and final invoices:\ni) All executed works will be certified by PEEIPL and payment released within 30 days after certification and submission of RA bills.\nii) Retention money @6% shall be withheld from each bill, released on successful completion after expiry of defect liability period.\niii) {vendor_name} shall submit proof/PF deposit Challan and other relevant documents with PEEIPL along with running account bills.\niv) At the time of submission of final certified invoice, {vendor_name} shall submit a declaration that there is no further claim under this LOI/Purchase order.\n\n{delivery_terms}"
    },
    {
        "number": "8",
        "title": "Taxes and Duties",
        "body": "8.1 BOQ rates are exclusive of all applicable taxes. {vendor_name} shall provide proof of GST registration, PAN etc. on award of contract. Any levy/penalty levied due to {vendor_name} to PEEIPL shall be recovered from {vendor_name}.\n\n8.2 Tax invoices are required for enabling PEEIPL to claim appropriate tax benefit.\n\n8.3 {vendor_name} shall be fully responsible for meeting all tax obligations and shall keep PEEIPL fully indemnified.\n\n8.4 TDS will be deducted at higher rate U/s 206AB, if {vendor_name} fails to file the ITR for last two years."
    },
    {
        "number": "9",
        "title": "Extra Items",
        "body": "i) If it is not similar item then rate will be derived from identical item of sub-contract.\nii) Rate for Extra item will be applicable from similar item of sub contract.\niii) If not possible to derive rate, it will be derived from the prevailing market rate plus 10% towards overhead and profit. CPWD norms will follow."
    },
    {
        "number": "10",
        "title": "Performance Guarantees",
        "body": "{warranty_terms}"
    },
    {
        "number": "11",
        "title": "Effective Date",
        "body": "Effective date of this Work/Scope shall be the date of this work/Scope. {vendor_name} shall proceed with the mobilization of personnel/workmen, materials, equipments, machinery etc. required for execution of the works assigned.\n\nRequired By: {required_by}"
    },
    {
        "number": "12",
        "title": "Reference Documents",
        "body": "The Contract between the Client and PEEIPL and the Standard Specifications and guidelines of Client/Consultant/PEEIPL which are forming part of the contract between {vendor_name} and PEEIPL and the owner, shall be the reference and guiding documents for all purposes."
    },
    {
        "number": "13",
        "title": "Sub-Contractor to Arrange Facilities at its Own Cost",
        "body": "{vendor_name} shall engage sufficient number of manpower/personnel/Engineer along with experienced and competent supervisors to ensure quality of work and smooth and uninterrupted progress of the works. {vendor_name} shall not be allowed to continue with the work in the absence of appointed supervisor.\n\n{vendor_name} shall ensure and arrange at its own cost the accommodation, transportation, boarding etc. for its manpower/personnel deployed for the execution of works."
    },
    {
        "number": "14",
        "title": "Manpower",
        "body": "{vendor_name} shall submit with PEEIPL, the proposed site organization to be set up for the execution of the work. A chart showing the manpower allocation and deployment to this job with specific job allocations shall be submitted by {vendor_name} to PEEIPL. {vendor_name} shall take prior approval from PEEIPL for any change in the allocation and deployment."
    },
    {
        "number": "15",
        "title": "Equipment and Machinery to be Deployed",
        "body": "{vendor_name} shall mobilize all required tools and machinery for execution of work. Mobilization Plan shall be prepared and monitored from time to time by {vendor_name}. {vendor_name} shall ensure to get fitness certificate for all its equipment and machinery from the Owner or the agency notified by the Owner."
    },
    {
        "number": "16",
        "title": "Licenses and Permits",
        "body": "{vendor_name} shall obtain at its own cost, all the licenses and permits required under the provisions of applicable Acts/Statutes, Regulations and rules for execution of various works under this WORK/Scope/contract.\n\n{vendor_name} shall obtain and keep valid license under the provisions of Contract Labour (Regulation and Abolition) Act, 1970 and other licenses wherever necessary.\n\n{vendor_name} shall submit to PEEIPL copies of all necessary licenses/permissions/permits before commencement of works."
    },
    {
        "number": "17",
        "title": "Insurance Cover/Policies",
        "body": "Before commencing execution of the Work, {vendor_name} shall insure against liability for loss of any material, equipment, machinery or physical damage, loss or injury arising out of the execution of the Works.\n\n(A) Contractor's All Risk Insurance Policy covering: entire WORK/Scope/Contract value, third party insurance, civil commotion/riots/war, earthquake, fire, and any other applicable insurance policy.\n\n(B) Policy to cover {vendor_name}'s liability under Workmen's Compensation Act 1923, Minimum Wages Act 1948, Contract Labour (Regulation and Abolition) Act 1970.\n\n(C) Insurance cover against damage or loss in respect of materials, equipment and/or work done. Limit of liability shall not be less than the value of such materials at any stage of the Contract."
    },
    {
        "number": "18",
        "title": "Sub-Contractor's Liability",
        "body": "{vendor_name} hereby assumes liability for and agrees to save PEEIPL harmless and indemnifies from every expense, liability or payment by reason of any injury (including death) to any person or damage to property suffered through any act or omission of {vendor_name}, his employees, workmen or from the conditions of the Site which is in the control of {vendor_name} for execution of the works."
    },
    {
        "number": "19",
        "title": "Compliances under Applicable Acts/Legislations, Rules and Regulations",
        "body": "(i) {vendor_name} shall abide by all Acts/statutes, Legislations/Rules and Regulations as applicable to the said WORK/Scope and ensure and pay wages to its workmen as per the Minimum Wages Act, 1948.\n(ii) {vendor_name} shall be entirely responsible for compliances of all applicable provisions under the Employees Provident Fund and Miscellaneous Provisions Act, 1952, Workman's Compensation Act, 1923 and other applicable Acts/statutes.\n(iii) {vendor_name} shall obtain comprehensive insurance cover for its entire manpower against any injury/death during execution of works.\n(iv) {vendor_name} shall indemnify and keep indemnified PEEIPL against all claims, liabilities and expenses arising out of default/breach of any statutory provisions.\n(v) {vendor_name} shall maintain statutory records viz. muster roll, payment register etc.\n(vi) {vendor_name} shall disburse wages in presence of authorized representative of PEEIPL.\n(vii) {vendor_name} shall submit with PEEIPL proof/challan pertaining to deposit of PF and other statutory payments."
    },
    {
        "number": "20",
        "title": "Health, Safety and Environment Related Regulations",
        "body": "(a) Healthy and hygienic Conditions: {vendor_name} shall ensure suitable welfare and hygiene arrangements at the site and shall follow applicable rules and regulations.\n\n(b) Safety of Site and Safety Equipments: {vendor_name} shall take full responsibility for the adequacy, stability and safety of all Site operations. {vendor_name} shall arrange sufficient helmets, safety boots/shoes and protective clothing for workmen.\n\n(c) Protection of Environment: {vendor_name} shall comply with all applicable environmental laws and regulations and shall ensure that the Site remains free from pollutants. Notwithstanding the above, {vendor_name} shall comply with all the directions and decisions of PEEIPL in this regard."
    },
    {
        "number": "21",
        "title": "Date of Commencement and Completion",
        "body": "Time being the essence of this contract, {vendor_name} shall ensure and be entirely responsible for completion of the entire works under this WORK/Scope within the agreed timeline from receipt of this LOI/work order and as per project schedule given by Project In-charge.\n\nRequired By: {required_by}"
    },
    {
        "number": "22",
        "title": "Liquidated Damages",
        "body": "Back-to-Back as per Agreement executed between PEEIPL and {vendor_name}.\n\n{penalty_clauses}"
    },
    {
        "number": "23",
        "title": "Professional Performance",
        "body": "{vendor_name} has warranted that it shall perform the WORK/Scope/Contract in a professional manner, using sound engineering principles, procedures and practices and with such care and diligence as are required by and in accordance with the standards of care customarily practiced by reputed and leading contractors. {vendor_name} represents that it has the required skills and capacity to perform the Services."
    },
    {
        "number": "24",
        "title": "Compliance to Specifications and Other Requirements",
        "body": "{vendor_name} shall comply with the standard specifications and other technical requirements for execution of the assigned work as defined and laid down by the Client/Consultant/PEEIPL. {vendor_name} shall also comply with Owner's/PEEIPL's inspection requirements and measurement instructions.\n\nAcceptance of any technical or specification deviation by {vendor_name} shall be subject to acceptance by PEEIPL. In case of non-acceptance, {vendor_name} shall execute the job without any deviations or extra time and cost implications to PEEIPL."
    },
    {
        "number": "25",
        "title": "Protection of Underground Utilities and Repair of Damages",
        "body": "PEEIPL shall provide all available details of underground utilities to {vendor_name}. {vendor_name} shall obtain plans and full details of all existing and planned underground utilities/services from the relevant Local Authorities. {vendor_name} shall be fully responsible for location and protection of all underground lines and structures.\n\nShould any damage occur, {vendor_name} shall immediately contact the concerned person/authority and repair work shall forthwith be carried out by {vendor_name} at its own expenses."
    },
    {
        "number": "26",
        "title": "Free Issue/Supply of Material",
        "body": "All free issue material, supplied to {vendor_name} by PEEIPL, shall be properly stored and handled by {vendor_name} and kept entirely separate for easy identification. {vendor_name} shall keep a proper record showing details of the materials issued from the storage area and the balance remaining available for use.\n\n{vendor_name} shall be solely responsible and liable for safe keeping and safe custody of all free issue material. Wastage limit of free issue materials shall be as per provision in the main Contract and CPWD norms.\n\nReconciliation will be done on monthly basis; additional wastage of material will be recovered from RA bills with 20% extra handling charges and applicable taxes."
    },
    {
        "number": "27",
        "title": "Subletting/Assignments",
        "body": "{vendor_name} shall not, without the prior written approval of PEEIPL, subject or assign to any other third party the whole or any portion of the work under this contract. If such approval is granted, {vendor_name} shall not be relieved of any of its obligations, duties and responsibility under this WORK/Scope/contract."
    },
    {
        "number": "28",
        "title": "Confidentiality",
        "body": "{vendor_name} understands and agrees to treat as strictly confidential all the technical data and information handed over by PEEIPL in terms of this WORK/Scope. {vendor_name} shall not disclose or reveal the technical data and information provided by PEEIPL to any third party except its employees, if essential and strictly on need-to-know basis.\n\nConfidentiality clause shall not be applicable in respect of information already in possession of either party prior to its disclosure."
    },
    {
        "number": "29",
        "title": "Defect Liability Period",
        "body": "Back-to-Back as per Agreement executed between PEEIPL and {vendor_name}.\n\n{warranty_terms}"
    },
    {
        "number": "30",
        "title": "Indemnification by the Sub-Contractor",
        "body": "{vendor_name} hereby agrees to indemnify and shall keep PEEIPL indemnified and harmless from and against any and all liabilities, losses, damages, costs, claims, actions, proceedings, expenses which may be suffered or incurred by PEEIPL as a result of any misrepresentation or breach of terms by {vendor_name} under this WORK/Scope."
    },
    {
        "number": "31",
        "title": "Supersession / Entire WORK/Scope/Contract",
        "body": "This WORK/Scope/Contract, including the Annexure(s) attached hereto, constitutes and represents the entire WORK/Scope/Contract between the parties and cancels and supersedes all prior understandings, letters, agreements, representations, statements, negotiations between the parties in respect of the matters dealt with herein."
    },
    {
        "number": "32",
        "title": "Amendments",
        "body": "No amendment, supplement, modification or clarification of this WORK/Scope/Contract shall be valid or binding unless set forth in writing and duly executed by the parties to this WORK/Scope/contract."
    },
    {
        "number": "33",
        "title": "Captions and Headings",
        "body": "Captions and Headings, as used herein, are for convenience of reference only and shall not be construed to limit or extend the language of the provisions to which such captions or Heading may refer in this WORK/Scope/Contract."
    },
    {
        "number": "34",
        "title": "Severability",
        "body": "If any provision of this WORK/Scope/contract is determined to be invalid or unenforceable in whole or in part, such invalidity or unenforceability shall attach only to such provision and the remaining provisions of the WORK/Scope/contract shall continue to remain in full force and effect."
    },
    {
        "number": "35",
        "title": "Force Majeure",
        "body": "The failure of a Party to fulfill any of its obligations under the WORK/Scope shall not be considered to be a breach or default insofar as such inability arises from an event of Force Majeure, i.e. fire, tempest, flood, earthquake, war, civil disturbances, change in government policies, violence of an army or mob or terrorist attack, caused not due to act/s or omission/s of the Party, provided that the Party:\n(a) has taken all reasonable precautions and due care; and\n(b) has informed the other Party as soon as possible about the occurrence of such an event.\n\nShould one or both the Parties be prevented from fulfilling their contractual obligations by a state of Force Majeure lasting continuously for a period of one month, both parties should consult with each other regarding future implementation of the WORK/Scope."
    },
    {
        "number": "36",
        "title": "PEEIPL's Right to Engage Another Sub-Contractor",
        "body": "In the event that {vendor_name} unjustifiably fails to complete the entire works or part of the works assigned within the agreed and specified period(s), PEEIPL shall have full right to engage other Sub contractor/deploy additional manpower/machinery and get the works executed and completed by such other Sub contractor at the sole cost and risk of {vendor_name}.\n\nFurther, {vendor_name} shall also be liable to pay to PEEIPL damages for such breach to the extent PEEIPL suffered the loss, without prejudice to any other rights or guarantees enforceable under this WORK/Scope/Contract."
    },
    {
        "number": "37",
        "title": "Suspension",
        "body": "i) {vendor_name} shall, if instructed in writing by PEEIPL, suspend the works or any part thereof for such period so ordered and shall not proceed until {vendor_name} shall have received a written instruction from PEEIPL to commence the works.\nii) Unless the reason of such suspension is the default of {vendor_name}, {vendor_name} shall be entitled to an adjustment of Time Schedule for that period.\niii) Upon suspension of the works, {vendor_name} shall exercise all reasonable efforts to preserve and safeguard the suspended works and continue to complete performance of the balance of the works, if applicable."
    },
    {
        "number": "38",
        "title": "Termination",
        "body": "PEEIPL shall have the right to terminate this WORK/Scope/contract by giving 7 (seven) days advance written notice to {vendor_name}, where {vendor_name} becomes bankrupt or where due to any act, deed or omission on the part of {vendor_name}, results in breach of any term and condition of this WORK/Scope or any default which being capable of cure has not been cured within ten days from the date of receipt of notice issued by PEEIPL to {vendor_name}."
    },
    {
        "number": "39",
        "title": "Address for Communication/Notice",
        "body": "All communications/notices between the parties shall be sent through Registered A/D Post/Courier service at the address of the parties:\n\ni) In case of communications to {vendor_name}:\nAttention: {vendor_contact}\n{delivery_address}\n\nii) In case of communications to PEEIPL:\nPassavant Energy & Environment India Pvt. Ltd.\nNavi Mumbai, India"
    },
    {
        "number": "40",
        "title": "General",
        "body": "a) {vendor_name} shall ensure that all its workmen deployed for execution of the work shall be in conformity with the applicable statutory provisions and laws enacted from time to time by the Government.\n\nb) PEEIPL and {vendor_name} have entered into this contract on Principal-to-Principal basis and nothing stated herein shall be deemed or construed as a partnership or as a joint venture or as an agency.\n\nc) Each party is and shall remain an Independent Party. None of the Party or any of its Affiliates shall be considered an agent of the other.\n\nd) Neither party shall be deemed to have waived any right under this WORK/Scope/contract, unless such party has delivered to the other party a written waiver signed by that party or its duly authorized signatory.\n\ne) Nothing in this WORK/Scope will preclude PEEIPL from having similar relationships with other Sub contractors."
    },
    {
        "number": "41",
        "title": "Dispute Mechanism",
        "body": "(a) Resolution/Settlement through mutual discussion and negotiation:\nIn the event of any dispute or difference arising out of or in connection with the WORK/Scope/contract, the Parties hereto shall at the first instance use their best efforts to settle such disputes or differences amicably by mutual discussion and negotiation.\n\n(b) Arbitration:\nIn case the amicable resolution or settlement is not reached within a period of 30 days, such dispute(s) or difference(s) shall be referred to a sole Arbitrator for settlement by way of arbitration in accordance with the provisions of the Arbitration and Conciliation Act 1996. The sole arbitrator shall be appointed by the mutual consent of both the parties. The decision of the Arbitrator shall be final and binding on both the Parties. The venue of such arbitration shall be at New Delhi."
    },
    {
        "number": "42",
        "title": "Jurisdiction",
        "body": "In case of any dispute arises between the Parties relating to the construction, meaning and operation of this WORK/Scope or breach thereof, the courts in Gurgaon alone shall have the Jurisdiction.\n\nThis LOI/work order is being issued and sent to {vendor_name} in duplicate; {vendor_name} is required to send to PEEIPL one copy of the LOI/work order together with all its attachments duly signed and stamped in token of their unconditional acceptance of the same. If signed copy is not received within 3 days of issuance of LOI/Work order, this LOI/Work order is deemed by PEEIPL to be accepted in its entirety by {vendor_name}.\n\n{special_conditions}"
    }
]
SUPPLY_ARTICLES  = [a.copy() for a in TECHNOLOGY_ARTICLES]

# ── Material PO – Annexure-II Special Conditions ──────────────────────────────
# These are stored as "articles" so the LOI editor UI works, but the download
# uses generate_material_po_pdf / generate_material_po_docx which renders the
# full PO document structure (PO header + Annexure-I items + Annexure-II + III).
MATERIAL_SPECIAL_CONDITIONS = [
    {
        "number": "1",
        "title": "PRICES",
        "body": (
            "PO value is INR {total_amount} ({total_amount_words}) on FOR site basis.\n"
            "The above accepted price is inclusive of Packaging & Forwarding, transit insurance "
            "and transportation to site, unloading PEEIPL Scope.\n"
            "The above accepted price is exclusive of GST and shall be paid extra as applicable.\n"
            "TDS will be deducted as per applicable laws.\n"
            "All payments made under this agreement shall be subject to reduction to reflect taxes "
            "or other charges required to be withheld by law.\n"
            "The rates given in the order is firm and not subject to any price escalation whatever "
            "may be the reason for the total period of execution of works under this purchase order "
            "including the extended period, if any and till the date of completion of the work and "
            "shall not be subject to any variation whatsoever."
        ),
    },
    {
        "number": "2",
        "title": "SCOPE OF WORK",
        "body": (
            "Vendor will finally be responsible for Supply of all the goods to start the Supply "
            "as complete unit complying with tender specification, client requirement & approved QAP, "
            "Datasheet & GA Drawing etc.\n\n"
            "Vendor should submit 1 Original + soft copy of all the MTCs, TCs, Inspection report, "
            "Internal inspection report & also other documents related to Engineering/design/QA/QC "
            "to be provided as a part of supply in file in our format provided during inspection "
            "stages & with dispatch clearance.\n\n"
            "Vendor should provide the one final set soft copy for our & end use evaluation before "
            "final submissions.\n\n"
            "O2 Visits at project sites shall be provided by supplier. Nos. of days shall be decided "
            "as per the site requirements. However, in case in additional visits required the same "
            "will be provided by SUPPLIER free of cost basis.\n\n"
            "Vendor will have to raise the inspection call before 7 days in advance before readiness "
            "of equipment. PEEIPL will arrange the inspection accordingly. Inspection of Entire order "
            "shall be completed in 1 visit, all charges (fooding, lodging, local travel etc) with "
            "respect the same will be taken care by KAY. Flight charges shall be borne by PEEIPL. "
            "TPI charges if any will be borne by PEEIPL.\n\n"
            "During approval, if any suggestion/change is suggested by Client, Vendor to accept the "
            "same without any additional charges to PEEIPL.\n\n"
            "{description}"
        ),
    },
    {
        "number": "3",
        "title": "TERMS OF PAYMENT",
        "body": (
            "10% payment against submission of ABG of equivalent amount after drawing approval.\n"
            "80% after delivery of material at site within 45 days from receipt of materials at site.\n"
            "10% payment will be released after commissioning and handing over or after maximum 180 days "
            "in case of delay in commissioning whichever is earlier. Pre-condition of such release will "
            "submission of PBG of equivalent amount valid till warrantee period:\n\n"
            "  \u2022 Original invoice\n"
            "  \u2022 Delivery Challan\n"
            "  \u2022 Packing List\n"
            "  \u2022 Original LF\n"
            "  \u2022 Warranty certificate\n"
            "  \u2022 Inspection report and Written Dispatch Clearance by PEEIPL.\n"
            "  \u2022 Any other document, BUYER may reasonably ask for.\n\n"
            "{delivery_terms}"
        ),
    },
    {
        "number": "4",
        "title": "DELIVERY",
        "body": (
            "Delivery shall be completed within 6 weeks from issuance of manufacturing clearance. "
            "Vendor shall expedite the delivery in 4 weeks from MFC.\n\n"
            "Required By: {required_by}\n"
            "Delivery Address: {delivery_address}"
        ),
    },
    {
        "number": "5",
        "title": "DISPATCH INSTRUCTIONS",
        "body": (
            "Seller shall dispatch the material through their own arrangement as per dispatched "
            "instructions/clearance given by Engineer-in-charge."
        ),
    },
    {
        "number": "6",
        "title": "WARRANTY",
        "body": (
            "Warranty for supplied material/equipment shall be 30 months from the date of supply "
            "or 24 months from the date of commissioning whichever is earlier.\n\n"
            "{warranty_terms}"
        ),
    },
    {
        "number": "7",
        "title": "DOCUMENTS",
        "body": (
            "Unless otherwise specifically changed/modified in this purchase order, all other terms "
            "and conditions of LOI No. dated 02.05.2026 shall be Integral Part of this Purchase Order."
        ),
    },
]

# Annexure-III General Conditions (clauses 1–34 from the PO images)
MATERIAL_GENERAL_CONDITIONS = [
    {
        "number": "1",
        "title": "DEFINITIONS",
        "body": (
            'These General Conditions of Purchase are hereinafter referred to as the "Conditions". '
            "In these Conditions, the following terms shall have the meanings set out below:\n\n"
            '"Contract": The set of contractual documents concerning the Supply and governing '
            "relations between the Seller and Buyer\n\n"
            '"Buyer": means PASSAVANT ENERGY & ENVIRONMENT INDIA PVT. LTD.\n\n'
            '"Seller": The Seller selected by the Buyer to perform the Contract.\n\n'
            '"Client": means MITL, the end user to whom the goods / services are being supplied by Buyer.\n\n'
            '"Consultant": Consultant selected by the client to provides expert technical advice to '
            "execute & complete the project.\n\n"
            '"Purchase Order/Contract": shall have same meaning.'
        ),
    },
    {
        "number": "2",
        "title": "CONTRACTUAL DOCUMENTS",
        "body": (
            "These Conditions shall prevail and shall override any general conditions of sale as well "
            "as any provision contained in or on the back of any invoices and/or in any other document "
            "issued by the Seller, and shall be applicable in so far as no legal, contractual or "
            "specific condition set out in the Contract stipulates otherwise.\n\n"
            "Changes, amendments, or modification of this agreement including its exhibits shall be "
            "valid only if they are in writing and signed by the officers of the parties.\n\n"
            "Failure to exercise part of any right under this Contract in one or more instances shall "
            "not constitute a waiver of those rights, waiver by one party of any of the rights "
            "established herein shall not be considered as a waiver of another right established herein.\n\n"
            "The Contract constitutes the entire agreement of the parties and supersedes all previous "
            "exchanges, obligations and agreements relating to the Supply."
        ),
    },
    {
        "number": "3",
        "title": "ACCEPTANCE",
        "body": (
            "The confirmation of purchase order shall constitute the contract and shall be sent by "
            "the Seller within 7 days of this order failing which it would be assumed that the order "
            "has been accepted by the Seller."
        ),
    },
    {
        "number": "4",
        "title": "COMMENCEMENT",
        "body": "The contract shall be operative from the date of the order.",
    },
    {
        "number": "5",
        "title": "QUALITY",
        "body": (
            "Goods supplied against order must in all respects conform to the descriptions given out "
            "of this order and drawings / specifications/ Approved QAP/samples approved and/or "
            "provided. Goods supplied against the order shall be perfect of material and workmanship "
            "and fully adaptable for the use intended and climatic conditions at the installation site.\n\n"
            "5.1 Quality control / audits / Inspection\n\n"
            "The Buyer shall be entitled, either itself or through its authorized representatives, "
            "to carry out quality controls and/or audits/Inspection at the Seller's / Subcontractor's "
            "production sites before or during the performance of the Contract. These quality controls "
            "and/or audits carried out by the Buyer shall not reduce the Seller's contractual liability "
            "in any way whatsoever, for example regarding the extent of the Seller's own required "
            "quality controls and further they shall not affect the Buyer's right to subsequently "
            "reject all or part of the Supply on delivery. The Seller shall provide to the Buyer all "
            "assistance required to carry out such quality controls and/or audits. In the event of "
            "rejection of Goods, due to defective design/material/manufacture/workmanship, the goods "
            "shall be rectified or new goods shall be offered by the Seller, for re-inspection as "
            "early as possible. The cost and expenses of such re-inspection shall be to Seller's account.\n\n"
            "Buyer reserves the right to claim free replacement/rectification cost for use of "
            "defective and sub-standard goods supplied by the Seller, irrespective of the fact whether "
            "the said goods were inspected and accepted by the Buyer.\n\n"
            "5.2 Traceability\n\n"
            "Upon written request by the Buyer, the Seller undertakes to furnish to the Buyer all "
            "information necessary to identify the origin, place and date of manufacture of the Supply; "
            "its components; the quality controls performed in relation to the Supply, the serial or "
            "batch numbers; and any other relevant information."
        ),
    },
    {
        "number": "6",
        "title": "PACKING – LABELING - MARKING",
        "body": (
            "The Seller shall be responsible for packing the Supplies. The packing shall be suitable "
            "for the Supply delivered and the means of transport used to deliver the Supply. Furthermore, "
            "the packing shall be in conformity with all applicable legislation, state of the art, and "
            "shall be adequate so as to prevent any damage to the Supply during transport, handling and "
            "storage at the place of delivery. The Supply shall be adequately labeled and packed in an "
            "appropriate manner and the parcels marked by the Seller in compliance with all applicable "
            "laws and moreover as provided for in the Contract. Special unloading/storage instructions "
            "must be mentioned in the packing list to be sent along with the supply. The Seller shall "
            "inform the Buyer in advance gross/net wt. of supply, Qty & dimensions often packages and "
            "the space requirements."
        ),
    },
    {
        "number": "7",
        "title": "ACCEPTANCE OF THE SUPPLY",
        "body": (
            "Acceptance of the Supply is subject to final approval of the Client's inspection regarding "
            "quality, quantity and specifications. Such inspection however does not absolve Seller of "
            "his responsibility for any defects in material/manufacturing/workmanship and guarantee & "
            "warranty conditions.\n\n"
            "If the Supply is expressly rejected, it shall be kept available at the Buyer's godown at "
            "the Seller's risk and expense. Client's inspection report and decision shall be considered "
            "as final and binding on the Seller.\n\n"
            "In the event of such rejection, and unless the Buyer decides otherwise in writing, the "
            "Supply shall, at the choice of the Buyer, either be repaired or be replaced not later than "
            "seven (7) calendar days following rejection by the Client. The Seller shall not raise any "
            "objection, including regarding its own manufacturing or delivery schedule, with regards to "
            "fulfilling the above obligation to repair or to replace.\n\n"
            "Seller must collect rejected goods within 7 days of intimation of rejection. In case Seller "
            "does not lift the rejected material within 7 days, Buyer reserves the right to dispose all "
            "the goods without any reference to the Seller.\n\n"
            "Payment made by the Buyer shall not constitute acceptance of the material.\n\n"
            "In case during further process of the supplied goods, any defect of whatsoever nature is "
            "observed, Buyer reserves the right to reject full/part quantity of such material event if "
            "in the first instance approved/paid by the Buyer.\n\n"
            "Quantity as received at the Buyer's godown/designated site address shall be treated as final."
        ),
    },
    {
        "number": "8",
        "title": "PRICES",
        "body": "The prices stated in the purchase order shall be firm till the execution of this order in full.",
    },
    {
        "number": "9",
        "title": "DELIVERY",
        "body": (
            "Delivery is the essence of the contract and supplies must be made as per the given schedule "
            "in the contract, Buyer can amend/postpone the delivery schedule on informing the same to "
            "Seller in writing. As soon as the Seller becomes aware of the fact that agreed deadlines "
            "cannot be upheld, he shall immediately inform Buyer in writing of the reasons and anticipated "
            "duration of the delay.\n\n"
            "In the event of premature deliveries without Buyer's consent, Buyer shall have the right to "
            "refuse acceptance. Payment tied to deadlines shall not become due before the agreed dates in "
            "the event of premature deliveries.\n\n"
            "The Seller shall not under any circumstances dispatch the Goods, until written clearance from "
            "Buyer is obtained, before dispatching the goods. Seller shall ensure that all the goods are "
            "dispatched in agreed sequence."
        ),
    },
    {
        "number": "10",
        "title": "VALIDITY",
        "body": "This purchase order shall remain valid till the completion of entire supplies and services under this order.",
    },
    {
        "number": "11",
        "title": "INVOICING AND PAYMENT",
        "body": (
            "The Buyer shall have no obligation/responsibility of payment where proof of delivery cannot "
            "be given satisfactorily.\n\n"
            "Invoices shall be prepared by the Seller in duplicate, in compliance with all applicable "
            "laws, and indeed the provisions of the Contract.\n\n"
            "The invoices shall be addressed to the Buyer at the address specified by the Buyer in the "
            "Order Form and shall quote the references of the Contract. All invoices shall be made out "
            "in the currency specified in the Contract.\n\n"
            "Details of batch no./identification nos. of the goods supplied must be mentioned on the invoice.\n\n"
            "The item code and description mentioned in the purchase order should be written on the "
            "invoice. No interest shall be paid on the overdue amount.\n\n"
            "Buyer reserves the right at all times to set off any amount which is due from the Buyer "
            "against any amount recoverable from the Seller in respect of any other transactions.\n\n"
            "Payment of the invoice shall not affect the Buyer's right to dispute in writing any "
            "unjustified charge."
        ),
    },
    {
        "number": "12",
        "title": "TEST CERTIFICATES",
        "body": "The Seller shall send relevant test certificates along with the supplies.",
    },
    {
        "number": "13",
        "title": "VENDOR DATA REQUIREMENT",
        "body": "Seller shall submit all drawings/data/documents/manuals in requisite nos. and form as agreed.",
    },
    {
        "number": "14",
        "title": "TRANSPORTATION",
        "body": (
            "Seller shall ensure to send the goods through nominated/approved transporters for destination "
            "as mentioned in the order. Seller shall bear all losses if goods are not shipped as per "
            "shipment instructions given in the order."
        ),
    },
    {
        "number": "15",
        "title": "WARRANTIES AND OTHER UNDERTAKINGS",
        "body": (
            "15.1 - Purpose\n\n"
            "The Seller shall, irrespective of the Buyer having any prior competence or knowledge, give "
            "the Buyer all necessary information, advice and warnings in relation to the nature and "
            "composition of the Supply. The Seller shall warn the Buyer about the risks related to the "
            "Supply, including but not limited to health and safety risks or concerns and any other "
            "hazardous risks.\n\n"
            "The Seller warrants that it has the full right to sell the Supply and that the Supply is "
            "free from any encumbrances, rights and privileges of any third party. The Seller warrants "
            "that the Supply corresponds to any description, specification and to any samples referred "
            "to in the Contract. The Seller further warrants that the Supply is fit for the purpose(s) "
            "and objective(s), specified by the Buyer and shall have no recourse to any purported lack "
            "of accuracy in the documents attached to the Contract.\n\n"
            "The Seller shall observe all laws, rules, as applicable in India provisions and highest "
            "professional standards as per relevant IS applicable to the Supply, in particular with "
            "regard to production, manufacture, repair, price definition and delivery, in order to "
            "ensure that the said Supply may be legally purchased, sold, transported and exported.\n\n"
            "Buyer shall under no circumstances be liable for any penalty/compensation that may be "
            "levied upon the Seller for noncompliance of various statutory rules & regulations. All "
            "Goods shall be supplied strictly in accordance with the specifications, drawings, "
            "data-sheets, other attachments and conditions stipulated under contract. All goods supplied "
            "under the contract (irrespective of whether engineering, design-data or other information "
            "has been furnished, reviewed or approved) are guaranteed by the Seller to be new and of "
            "the best quality of their respective kinds and shall be free from any in cumbrance faulty "
            "design, materials, manufacture and workmanship and to be of sufficient size and capacity "
            "and of proper materials so as to fulfil in all respects the intended purpose specified by "
            "the Buyer.\n\n"
            "15.2 - Duration and scope\n\n"
            "The Seller warrants for a period as specified in the order, from the date of the Buyer's "
            "acceptance of delivery, that the Supply is free of manufacturing defects, faults, "
            "contamination and abnormal wear & tear of whatsoever nature. Should the Seller breach this "
            "warranty, the Buyer shall at its option either require the Seller to repair or to replace "
            "the Supply or if the Seller does not do so within mutually agreed period as of the date of "
            "the request by the Buyer, take appropriate measures to arrange the same itself or through "
            "a third party. In any event, the Seller shall bear all costs of any replacement and repair "
            "of the Supply, including but not limited to travel expenses, the costs of returning the "
            "Supply to the factory and any parts and labour, but without prejudice to any other rights "
            "or remedies the Buyer may have.\n\n"
            "Any replacement or repair of the Supply under warranty shall give rise to a new warranty "
            "for a minimum period of Twelve (12) months from the date of the Buyer's acceptance of "
            "delivery applicable for the repaired or replaced Supply.\n\n"
            "The Seller remains bound by all applicable statutory warranties and product liability "
            "guarantees as well as for hidden defects.\n\n"
            "15.3 - Availability of spare parts\n\n"
            "Not applicable."
        ),
    },
    {
        "number": "16",
        "title": "COMPLETENESS CLAUSE",
        "body": (
            "Except for any item specifically excluded, any material as per defined & mutually agreed "
            "scope, equipment and services, which may not be mentioned in the specifications or drawings, "
            "but which are implied or necessary for completion of the supplies under the contract, are "
            "to be provided by the Seller, without any extra cost to the Buyer and the equipment shall "
            "be complete in all respects to ensure that the same is capable of sustained performance as "
            "specified including meeting the specified performance test guarantees.\n\n"
            "In case the Buyer is required to incur any expenditure for completing the equipment, "
            "including for additions, alterations or modifications, due to reasons attributable to the "
            "Seller, this amount shall be to the account of the Seller. The Buyer also retains the right "
            "to reject the incomplete equipment and recover the amounts paid to the Seller, as well as "
            "procure the appropriate equipment at the Seller's cost and risk."
        ),
    },
    {
        "number": "17",
        "title": "HEALTH, SAFETY AND THE ENVIRONMENT",
        "body": (
            "When delivering the Supply to a site designated by the Buyer, the Seller shall observe and "
            "ensure that all of its employees, representatives or sub-contractors comply with all rules "
            "and regulations in force on the site designated by the Buyer regarding health, safety, "
            "working conditions and the environment as well as all applicable legislation and regulations.\n\n"
            "If the Seller, its Subcontractors, employees or representatives do not comply with any of "
            "these obligations, the Buyer reserves the right to refuse the Seller and/or any of its "
            "employees or sub-contractors' access to or continued presence at the site. All consequences "
            "of non-compliance with any of these obligations, and the denial of access to or of continued "
            "presence on the site, including any costs incurred by the Seller as a result of such denial, "
            "shall solely be borne by the Seller."
        ),
    },
    {
        "number": "18",
        "title": "ASSIGNMENTS OF ORDER / SUB-CONTRACTORS",
        "body": (
            "18.1 – Assignment and Change of Control\n\n"
            "The Seller shall not assign the Contract to any third party, in whole or in part, without "
            "the prior written consent of the Buyer. The Buyer shall be entitled to assign all or part "
            "of the Contract subject to prior written information regarding such assignment being sent "
            "to the Seller.\n\n"
            "18.2 - Sub-Contractors\n\n"
            "Under no circumstances shall the manufacturing of the Supply to be performed in accordance "
            "with specifications of the Buyer and the operations associated with such performance under "
            "the Contract be subcontracted or entrusted to another person or company by the Seller "
            "without the prior written consent of the Buyer and the same shall only be so subcontracted "
            "subject to the condition that the Seller and the said third party both comply with all "
            "statutory provisions applicable. Seller shall have to get its list of subcontractors "
            "approved by the Buyer, before placing sub-orders.\n\n"
            "In all cases, the Seller shall remain solely liable for the proper performance of the "
            "entire Contract. The Seller shall indemnify and hold the Buyer harmless from and against "
            "any and all claims by the Seller's employees or sub-contractors or by the personnel of "
            "said sub-contractors."
        ),
    },
    {
        "number": "19",
        "title": "STANDARD ELEMENTS",
        "body": (
            "To the extent that the Supply contains standard elements (including but not limited to "
            "plans, manuals or brochures, documents, standard software included in or incidental to the "
            "Supply) subject to intellectual property rights and handed over by the Seller to the Buyer "
            "for utilizing the Supply, the Seller grants to the Buyer and to the third parties acting "
            "on behalf of or for the benefit of the Buyer a personal and non-exclusive right to use, "
            "reproduce, represent, translate and adapt the said standard elements for its own purposes. "
            "This right shall be granted without additional cost for the entire duration of the "
            "applicable intellectual property right protection, and where permitted by law, in "
            "perpetuity, for all countries and for all media.\n\n"
            "In case of transfer by the Buyer of the Supply to a third party, the above right to use on "
            "standard elements shall be transferable by the Buyer to such third party without "
            "necessitating specific consent to such transfer from the Seller and without additional cost "
            "to the Buyer."
        ),
    },
    {
        "number": "20",
        "title": "INFRINGEMENT OF INTELLECTUAL PROPERTY RIGHTS",
        "body": (
            "The Seller declares that all intellectual property rights relating to the Supply are "
            "either the exclusive property of the Seller or subject to a license granted to the Seller "
            "by a third party holding these rights, under conditions allowing the Buyer to freely use "
            "and/or transfer the Supply.\n\n"
            "Consequently, the Seller shall indemnify and hold harmless the Buyer against any and all "
            "claims, costs, damages, expenses or legal action by third parties in connection with any "
            "actual or alleged infringement of intellectual property rights in relation to the Supply. "
            "If in the opinion of the Seller the Supply carries a risk of legal claim or legal action, "
            "the Seller shall take all steps necessary to ensure that such risk of infringement is "
            "eliminated. If an allegation is made that the Buyer may not use the Supply without "
            "infringing a third party's intellectual property right, the Seller shall, at its own cost "
            "and at the sole option of the Buyer, either replace or modify the Supply in respect of "
            "which such allegation is made, such a way so as to ensure that the infringement of "
            "intellectual property rights no longer exist and at all times in accordance with its "
            "contractual obligations. Such replacement or modification shall be performed within time "
            "scales compatible with the requirements of the Buyer. Failing such replacement or "
            "modification, the Seller shall refund to the Buyer the price of the Supplies. The above "
            "provisions do not affect the Buyer's right to claim further damages from the Seller."
        ),
    },
    {
        "number": "21",
        "title": "CODE OF CONDUCT",
        "body": (
            "The Seller shall acquaint itself and comply with the principles of Buyer's Code of Conduct. "
            "The Seller shall also ensure that all of its sub-contractors comply with the same."
        ),
    },
    {
        "number": "22",
        "title": "CONFIDENTIALITY",
        "body": (
            "Any document or information provided by the Buyer to the Seller in connection with the "
            "performance of the Contract, and all elements (including all statements, studies, and other "
            "documents) prepared by Seller for the performance of the Contract shall be treated as "
            "strictly confidential by the Seller. Furthermore, any document and information the Seller "
            "may acquire in connection with the performance of the Contract (including but not limited "
            "to any documents or information concerning the Buyer's organization, business activities, "
            "or financial results) shall be treated as strictly confidential by the Seller.\n\n"
            "The Seller shall only be entitled to use the documents, information and/or elements "
            "referred to above for the performance of the Contract and shall not disclose the same to "
            "any third party or to any staff or agent of the Seller other than to those who are involved "
            "in the performance of the Contract. This paragraph shall not apply to the extent disclosure "
            "is mandatory by virtue of any statutory, accounting or regulatory obligations.\n\n"
            "The Seller undertakes to comply with these obligations of non-use and of confidentiality "
            "and shall cause that its employees, agents and sub-contractors do the same, throughout the "
            "duration of the Contract and for a further period of five (5) years after the termination "
            "of the Contract.\n\n"
            "The Seller shall however not be liable for the disclosure of information to the extent "
            "that such information is in the public domain or has been legitimately obtained from other sources.\n\n"
            "The Seller shall at the expiry of the term of the Contract, return to the Buyer every "
            "document and data and every copy of the same, which may have been kept in connection with "
            "the performance of the Contract."
        ),
    },
    {
        "number": "23",
        "title": "REFERENCE TO THE BUYER'S BRANDS AND TRADE NAMES",
        "body": (
            "The Seller shall not be entitled to refer to the company names, trademarks and or logos "
            "of the Buyer's Group without the prior written consent of the Buyer."
        ),
    },
    {
        "number": "24",
        "title": "ORDER OF PRECEDENCE",
        "body": (
            "In case of any discrepancy between the conditions stated in the General terms & conditions "
            "and those specifically mentioned in the contract, the conditions mentioned in the contract "
            "shall supersede."
        ),
    },
    {
        "number": "25",
        "title": "FORCE MAJEURE",
        "body": (
            "Neither party hereto shall be liable to the other party if the performance of its "
            "obligations is delayed by conditions constituting force majeure under this clause. Force "
            "majeure shall only mean and include compliance with statutory laws, regulations or orders "
            "of the Central, State or local Government, Statutory Bodies, war and war-like conditions "
            "(irrespective of whether war has been declared or not), acts of civil and military "
            "authority, fires, floods, embargoes, sabotage, revolts, rebellion and strikes and "
            "lock-outs.\n\n"
            "PROVIDED THAT the strikes and lock-outs which do not exceed continuous period of fifteen "
            "days shall not be a force majeure condition, for the purpose of this clause and PROVIDED "
            "FURTHER THAT none of the conditions mentioned in this clause shall be a force majeure "
            "condition unless such conditions or conditions actually affect the performance of the "
            "obligations of either party. It is expressly agreed that power cuts shall not be considered "
            "as force majeure condition.\n\n"
            "Should however, any force majeure condition herein mentioned occur and continue for a "
            "period exceeding 15 days the parties hereto undertake to sit together and devise means "
            "for expeditious and proper performance of the obligations of the parties under this order."
        ),
    },
    {
        "number": "26",
        "title": "QUANTITY VARIATIONS",
        "body": (
            "Quantities are fixed & final and any variation i.e. increases in the quantity other than "
            "the specified in the purchase order shall not be considered for the payment."
        ),
    },
    {
        "number": "27",
        "title": "SELLER'S DEFAULT",
        "body": (
            "The Seller is considered in default in case any or all of the following events:-\n\n"
            "27.1.1 Fails to proceed with the purchase order with due diligence after being required "
            "to do so in writing.\n"
            "27.1.2 Fails to execute promptly his obligations in accordance with the Purchase Order "
            "after being requested in writing by the Buyer.\n"
            "27.1.3 Fails to remove defective materials or make good defective work after being "
            "directed in writing to do so.\n"
            "27.1.4 Commits an act of bankruptcy or goes into liquidation.\n"
            "27.1.5 Fails to successfully execute, deliver and complete the supply within stipulated time\n"
            "27.1.6 The Seller causes unreasonable delay to the project due to lack of materials or "
            "any reason whatsoever and in such event, the buyer shall back charge seller & recover such "
            "loss from seller that caused loss to the buyer.\n\n"
            "In the event of any or all of the above the buyer may, at his own option, terminate the "
            "Seller's Contract and: order the goods and services himself, or\n"
            "Employ the services of any other Seller, individual or otherwise to complete the delivery "
            "of goods and services.\n"
            "Upon such termination the buyer will short close the order at their end."
        ),
    },
    {
        "number": "28",
        "title": "DELAY",
        "body": (
            "In case of any of the following events:-\n\n"
            "  \u2022 The Seller delays delivery of products and any associated service as stipulated in the Contract.\n"
            "  \u2022 The Seller delivers goods partially thereby delay the buyer's schedule.\n"
            "  \u2022 The Seller delivers goods not in accordance with the Purchase Order/supply agreement "
            "causing delays to buyer Schedule of works.\n"
            "  \u2022 The Seller delivers unapproved goods to site and delays replacement of the approved "
            "materials after receiving written instructions to do so causing delays to the buyer's "
            "schedule of works.\n\n"
            "In case, the seller performance is delayed due to any delay on client's side, the seller "
            "shall be given due extension of time without any cost implication. The seller shall "
            "intimate client regarding such delay in writing.\n\n"
            "The Seller is liable for the one or more of the following actions:-\n"
            "  \u2022 The buyer is entitled to pursue the outstanding materials (not delivered portion) and/or "
            "replacement of materials or not approved materials in any means possible including freight etc.\n"
            "  \u2022 The buyer is entitled to seek replacement for any damages supplies (in transit or receipt)."
        ),
    },
    {
        "number": "29",
        "title": "LIQUIDATED DAMAGES",
        "body": (
            "If the Seller fails to successfully execute, deliver and complete the supply within "
            "stipulated time as mentioned in the contract, the seller shall be liable to pay to buyer "
            "Liquidated damages, a sum equivalent to 0.5% (Zero Point Five Per Cent) of the Ex-work "
            "contract value for each week of delay or any part thereof. However, total amount of "
            "liquidated damage for delay in completion of contract shall be subject to a maximum of "
            "5% of undelivered portion.\n\n"
            "{penalty_clauses}"
        ),
    },
    {
        "number": "30",
        "title": "SUSPENSION OF WORKS",
        "body": (
            "BY BUYER:\n\n"
            "The buyer may by written notice to the Seller suspend for a specified period, in whole "
            "or in part, payments to the Seller and/or the Seller's obligation to continue to perform "
            "the Works under this Contract, if in the buyer's sole discretion:\n\n"
            "(a) any conditions arise which interfere, or threaten to interfere with the successful "
            "execution of the supply works or the accomplishment of the purpose thereof, or\n\n"
            "(b) the Seller shall have failed, in whole or in part, to perform any of the terms and "
            "conditions of this Contract.\n\n"
            "The term of this Contract may be extended by the buyer for a period equal to any period "
            "of suspension, taking into account any special conditions which may cause the additional "
            "time for completion of the Works to be different from the period of suspension.\n\n"
            "BY SELLER:\n\n"
            "Seller may suspend any part or pull of the contract after issuing a written notice to "
            "buyer in case of any breach of the contract committed by buyer. Seller shall resume the "
            "work immediately on the buyer remedy the defect or breach of the contract."
        ),
    },
    {
        "number": "31",
        "title": "TERMINATION",
        "body": (
            "The Buyer reserves the right to terminate the Contract in case of breach, without "
            "compensation to the seller towards the following defaults and Article 27 & 28; if the Seller:\n\n"
            "is neglecting to perform his obligations or is not seriously carrying out the scope of "
            "work, or becomes bankrupt or insolvent\n"
            "Defaults in meeting the contract requirement which will jeopardize the project "
            "Completion plan.\n"
            "using a substandard or non-approved materials for the work\n"
            "Fails to comply with notice from the buyer. Abandons or repudiates the contract\n"
            "Sub contracts the whole or part of the contract or assigns the contract without approval "
            "of the buyer.\n"
            "Persistently disregards the instructions of the buyer or his authorized representative "
            "or contravenes any provision of the contract or fails to adhere to the programme of work\n"
            "Indulges in corrupt or fraudulent practices and practices against law of the land etc.\n\n"
            "The buyer shall be entitled to retain any balance, which may otherwise be due to the "
            "Seller and shall use this money to make the payment for executing the said part of the "
            "balance Work or of completing the Works as the case may be. In case, the retained money "
            "is not sufficient for completing the of the balance Work or of completing the Works as "
            "the case may be, the buyer reserves the right to recover this additional amount from any "
            "other payment due to the seller or from encashment of performance bank guarantee."
        ),
    },
    {
        "number": "32",
        "title": "LIMITATION OF LIABILITIES",
        "body": (
            "The aggregate liability of the Seller with respect to all claims arising out of or in "
            "connection with performance or non-performance of this Contract whether in contract, "
            "warranty, tort or otherwise shall not exceed the Contract Price, provided that this "
            "limitation shall not apply in case of negligence, wilful Misconduct or liabilities "
            "arising out of indemnity provisions in this Contract."
        ),
    },
    {
        "number": "33",
        "title": "ARBITRATIONS",
        "body": (
            "In case any dispute relating to the terms and conditions of this Contract or the "
            "interpretation thereof arises between the parties, the same shall promptly and in good "
            "faith be negotiated with a view of its amicable resolution and settlement. In the event "
            "no amicable resolution or settlement is reached within a period of 30 days from the day "
            "on which the dispute(s) or difference(s) arose, such dispute(s) or difference(s) shall "
            "be referred to and settled by the arbitration. In the event of dispute being referred "
            "for Arbitration three arbitrators shall be appointed. One each by the Buyer and Seller "
            "and third arbitrator shall be jointly appointed by the two Arbitrators.\n\n"
            "The existence of any dispute(s) or difference(s) or the initiation to continuance of "
            "the arbitration proceedings shall not permit the parties to postpone or delay the "
            "performance by the parties of their respective obligations pursuant to this Contract.\n"
            "The place of arbitration shall be Delhi.\n"
            "The decision and award resulting from such arbitration shall be final and binding on the Parties."
        ),
    },
    {
        "number": "34",
        "title": "JURISDICTIONS",
        "body": (
            "All suits, legal proceedings and arbitration award under this Contract shall be filed, "
            "entertained and decided in the Court of Delhi and the Courts in Gurgaon Haryana shall "
            "have the exclusive jurisdiction over all such disputes/claims.\n\n"
            "For Passavant Energy & Environment India Private Limited.\n\n"
            "(Authorized Signatory)"
        ),
    },
]

# Combined for the LOI editor UI (shown as single flat list)
MATERIAL_ARTICLES = MATERIAL_SPECIAL_CONDITIONS + MATERIAL_GENERAL_CONDITIONS

LOI_TEMPLATES = {
    "technology": TECHNOLOGY_ARTICLES,
    "service":    SERVICE_ARTICLES,
    "supply":     SUPPLY_ARTICLES,
    "material":   MATERIAL_ARTICLES,
}


def _amount_to_words(amount: float) -> str:
    """Convert a numeric INR amount to words (simplified, Indian system)."""
    try:
        amount = int(amount)
        if amount <= 0:
            return "Zero"
        crore  = amount // 10_000_000
        lakh   = (amount % 10_000_000) // 100_000
        thous  = (amount % 100_000) // 1000
        remain = amount % 1000
        parts = []
        ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight",
                "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
                "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty",
                "Sixty", "Seventy", "Eighty", "Ninety"]

        def below_100(n):
            if n < 20:
                return ones[n]
            return (tens[n // 10] + (" " + ones[n % 10] if n % 10 else "")).strip()

        def below_1000(n):
            if n >= 100:
                return ones[n // 100] + " Hundred" + (" " + below_100(n % 100) if n % 100 else "")
            return below_100(n)

        if crore:
            parts.append(below_1000(crore) + " Crore")
        if lakh:
            parts.append(below_100(lakh) + " Lakh")
        if thous:
            parts.append(below_100(thous) + " Thousand")
        if remain:
            parts.append(below_1000(remain))
        return " ".join(parts) + " only"
    except Exception:
        return ""


class _SafeFormatDict(dict):
    """dict for str.format_map that leaves unknown {placeholders} untouched."""
    def __missing__(self, key):
        return "{" + key + "}"


def _safe_fill(text: str, variables: dict) -> str:
    """Fill {placeholders} in text; never raise on stray braces or unknown keys."""
    if not text:
        return text or ""
    try:
        return text.format_map(_SafeFormatDict(variables))
    except Exception:
        # Unbalanced braces etc. — return the text as-is rather than failing
        return text


class LOIService:

    @staticmethod
    def _build_variables(po_data: dict[str, Any]) -> dict[str, str]:
        """Build the placeholder -> value map used to fill article bodies."""
        total = float(po_data.get("total_amount", 0))
        return {
            "vendor_name":        po_data.get("vendor_name", "[VENDOR NAME]"),
            "vendor_contact":     po_data.get("vendor_contact", "[VENDOR CONTACT]"),
            "po_number":          po_data.get("po_number", "[PO NUMBER]"),
            "total_amount":       f"INR {total:,.2f}",
            "total_amount_words": _amount_to_words(total),
            "description":        po_data.get("description", "[DESCRIPTION]"),
            "delivery_address":   po_data.get("delivery_address", "[DELIVERY ADDRESS]"),
            "required_by":        po_data.get("required_by", "[DATE]"),
            "penalty_clauses":    po_data.get("penalty_clauses") or "As per standard terms.",
            "delivery_terms":     po_data.get("delivery_terms") or "Within agreed timeline from the date of PO approval.",
            "warranty_terms":     po_data.get("warranty_terms") or "As per standard warranty terms.",
            "special_conditions": po_data.get("special_conditions") or "None.",
            "site_name":          po_data.get("site_name", "[SITE NAME]"),
        }

    @staticmethod

    def get_template_articles(po_type: str) -> list[dict]:
        """Return a copy of the articles for the given PO type."""
        return [
            {"number": a["number"], "title": a["title"], "body": a["body"]}
            for a in LOI_TEMPLATES.get(po_type, TECHNOLOGY_ARTICLES)
        ]

    @staticmethod
    def fill_articles(po_type: str, po_data: dict[str, Any]) -> list[dict]:
        """
        Fill placeholder variables in articles with actual PO data.
        Returns list of {number, title, body} dicts with placeholders replaced.
        """
        variables = LOIService._build_variables(po_data)
        filled = []
        for article in LOIService.get_template_articles(po_type):
            filled.append({
                "number": article["number"],
                "title":  article["title"],
                "body":   _safe_fill(article["body"], variables),
            })
        return filled

    @staticmethod
    def resolve_articles(
        po_type: str,
        stored_json: str | None,
        po_data: dict[str, Any],
    ) -> list[dict]:
        """
        Resolve the final article list for a PO.

        If the PO has a customized article set (stored_json, saved from the PO
        creation form), use it; otherwise fall back to the default template.

        stored_json is a JSON list of entries:
          {"src": "tpl", "idx": <template index>, "title": "...", "body": null}
            -> default template article kept; body null means "use the template
               body" (so placeholders are filled with real PO data here).
               A non-null body means the user edited the text on the form.
          {"src": "custom", "title": "...", "body": "..."}
            -> article added by the user on the form.

        Articles are renumbered sequentially 1..N in stored order.
        """
        import json as _json

        if not stored_json:
            return LOIService.fill_articles(po_type, po_data)
        try:
            entries = _json.loads(stored_json)
            if not isinstance(entries, list):
                raise ValueError("loi_articles is not a list")
        except Exception:
            return LOIService.fill_articles(po_type, po_data)

        tpl_articles = LOI_TEMPLATES.get(po_type, TECHNOLOGY_ARTICLES)
        variables = LOIService._build_variables(po_data)

        out: list[dict] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            src = entry.get("src")
            if src not in ("tpl", "custom"):
                continue
            title = str(entry.get("title") or "").strip()
            body = entry.get("body")
            if src == "tpl":
                idx = entry.get("idx")
                tpl = (
                    tpl_articles[idx]
                    if isinstance(idx, int) and 0 <= idx < len(tpl_articles)
                    else None
                )
                if tpl is None and body is None:
                    # Dangling template reference (e.g. template shrank) — skip
                    continue
                if tpl:
                    if body is None:
                        body = tpl["body"]
                    if not title:
                        title = tpl["title"]
            if body is None:
                body = ""
            out.append({
                "number": str(len(out) + 1),
                "title":  title or "UNTITLED",
                "body":   _safe_fill(str(body), variables),
            })

        return out or LOIService.fill_articles(po_type, po_data)

    @staticmethod
    def generate_pdf(
        po_data: dict[str, Any],
        articles: list[dict],
    ) -> bytes:
        """
        Generate a formatted PDF LOI document using fpdf2.
        fpdf2 is already in requirements.txt (fpdf2==2.8.7).
        """
        try:
            from fpdf import FPDF
        except ImportError:
            raise RuntimeError("fpdf2 is not installed. Run: pip install fpdf2 --break-system-packages")

        class LOI_PDF(FPDF):
            def header(self):
                self.set_font("Helvetica", "B", 10)
                self.set_fill_color(240, 240, 240)
                self.cell(0, 8, "M/S. PASSAVANT ENERGY & ENVIRONMENT INDIA PVT. LTD.", border=1, fill=True, ln=True, align="C")
                self.set_font("Helvetica", "", 9)
                self.cell(0, 5, "Navi Mumbai, India", ln=True, align="C")
                self.ln(3)

            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 8)
                self.cell(0, 10, f"Page {self.page_no()}", align="C")

        pdf = LOI_PDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_margins(20, 25, 20)
        pdf.add_page()

        vendor_name   = po_data.get("vendor_name", "[VENDOR NAME]")
        po_number     = po_data.get("po_number", "DRAFT")
        site_name     = po_data.get("site_name", "[SITE]")
        description   = po_data.get("description", "[DESCRIPTION]")
        vendor_contact = po_data.get("vendor_contact", "[CONTACT]")

        # PO meta
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, f"PO Number: {po_number}  |  Site: {site_name}", ln=True)
        pdf.ln(2)

        # Attn
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, f"Kind Attn: {vendor_contact}", ln=True)
        pdf.ln(2)

        # Subject
        pdf.set_font("Helvetica", "B", 10)
        subj = f"Subject: LOI for {description}"
        pdf.multi_cell(0, 6, subj)
        pdf.ln(3)

        # Intro
        pdf.set_font("Helvetica", "", 10)
        intro = (
            f"The contractor/buyer, Passavant Energy & Environment India Pvt. Ltd. (PEEIPL), "
            f"is pleased to issue this Letter of Intent to M/s {vendor_name} (the Seller) "
            f"subject to the terms and conditions set forth below."
        )
        pdf.multi_cell(0, 5, intro)
        pdf.ln(4)

        # Articles
        for article in articles:
            # Article heading
            pdf.set_font("Helvetica", "B", 10)
            heading = f"ARTICLE {article['number']} - {article['title']}"
            pdf.multi_cell(0, 6, heading)
            pdf.ln(1)

            # Article body - handle encoding issues gracefully
            pdf.set_font("Helvetica", "", 10)
            body = article.get("body", "")
            # fpdf2 handles UTF-8 but rupee symbol needs fallback
            body = body.replace("INR ", "INR ")
            try:
                pdf.multi_cell(0, 5, body)
            except Exception:
                pdf.multi_cell(0, 5, body.encode("latin-1", errors="replace").decode("latin-1"))
            pdf.ln(4)

        # Signing section
        pdf.ln(6)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 5, "Please sign and return a copy of this LOI as acceptance of the above terms.", ln=True)
        pdf.ln(8)

        # Two-column signing
        col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / 2
        y_before = pdf.get_y()
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(col_w, 6, "For Passavant Energy & Environment India Pvt. Ltd.")
        pdf.cell(col_w, 6, f"For {vendor_name}", ln=True)
        pdf.ln(14)
        pdf.set_draw_color(0, 0, 0)
        x = pdf.l_margin
        pdf.line(x, pdf.get_y(), x + col_w - 5, pdf.get_y())
        pdf.line(x + col_w + 5, pdf.get_y(), x + col_w * 2, pdf.get_y())
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(col_w, 5, "Authorised Signatory")
        pdf.cell(col_w, 5, "Authorised Signatory", ln=True)

        from io import BytesIO
        buf = BytesIO()
        pdf.output(buf)
        return buf.getvalue()

    @staticmethod
    def generate_docx(
        po_data: dict[str, Any],
        articles: list[dict],
    ) -> bytes:
        """
        Generate a formatted Word DOCX LOI document.
        Uses python-docx (must be installed: pip install python-docx).
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed. "
                "Run: pip install python-docx --break-system-packages"
            )

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Header - buyer info
        header_text = (
            "M/S. PASSAVANT ENERGY & ENVIRONMENT INDIA PVT. LTD.\n"
            "Navi Mumbai, India"
        )
        h = doc.add_paragraph(header_text)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.bold = True
            run.font.size = Pt(11)

        doc.add_paragraph()

        # Attn & Subject
        vendor_name = po_data.get("vendor_name", "[VENDOR NAME]")
        doc.add_paragraph(f"Kind Attn: {po_data.get('vendor_contact', '[CONTACT]')}")

        subj = doc.add_paragraph()
        subj_run = subj.add_run(f"Subject: LOI for {po_data.get('description', '[DESCRIPTION]')}")
        subj_run.bold = True

        doc.add_paragraph()

        # Intro paragraph
        intro = (
            f"The contractor/buyer, Passavant Energy & Environment India Pvt. Ltd. (PEEIPL), "
            f"desiring that certain works pertaining to the project as described herein, "
            f"has entered into discussions with M/s {vendor_name} (the Seller). "
            f"Purchase Order: {po_data.get('po_number', '[PO NUMBER]')} | "
            f"Site: {po_data.get('site_name', '[SITE]')}"
        )
        doc.add_paragraph(intro)
        doc.add_paragraph()

        # Articles
        for article in articles:
            # Article heading
            heading = doc.add_paragraph()
            run = heading.add_run(
                f"ARTICLE {article['number']} - {article['title']}"
            )
            run.bold = True
            run.font.size = Pt(11)

            # Article body - split on newlines for readability
            for line in article["body"].split("\n"):
                p = doc.add_paragraph(line)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(2)
            doc.add_paragraph()

        # Closing
        doc.add_paragraph(
            "Please sign and return a copy of this Letter of Intent as your acceptance "
            "of the above terms and conditions."
        )
        doc.add_paragraph()
        doc.add_paragraph("For Passavant Energy & Environment India Pvt. Ltd.")
        doc.add_paragraph("\n\nAuthorised Signatory")
        doc.add_paragraph()
        doc.add_paragraph(f"For {vendor_name}")
        doc.add_paragraph("\n\nAuthorised Signatory")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _build_html(po_data: dict[str, Any], articles: list[dict]) -> str:
        """Build the HTML that WeasyPrint renders to PDF."""
        vendor_name  = po_data.get("vendor_name", "[VENDOR NAME]")
        po_number    = po_data.get("po_number", "[PO NUMBER]")
        site_name    = po_data.get("site_name", "[SITE]")
        description  = po_data.get("description", "[DESCRIPTION]")
        vendor_contact = po_data.get("vendor_contact", "[CONTACT]")

        articles_html = ""
        for article in articles:
            body_html = article["body"].replace("\n\n", "</p><p>").replace("\n", "<br>")
            articles_html += f"""
            <div class="article">
                <div class="article-title">ARTICLE {article['number']} - {article['title']}</div>
                <p>{body_html}</p>
            </div>
            """

        return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  @page {{
    margin: 2cm 2.5cm;
    size: A4;
  }}
  body {{
    font-family: 'Times New Roman', Times, serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #000;
  }}
  .header-box {{
    border: 1px solid #000;
    padding: 14px 18px;
    margin-bottom: 18px;
  }}
  .company-name {{
    font-weight: bold;
    font-size: 12pt;
    text-transform: uppercase;
    margin-bottom: 4px;
  }}
  .company-address {{
    font-size: 10pt;
  }}
  .attn {{
    margin: 12px 0 4px;
    font-weight: bold;
  }}
  .subject {{
    font-weight: bold;
    font-size: 11pt;
    margin: 8px 0 14px;
  }}
  .intro {{
    margin-bottom: 14px;
    text-align: justify;
  }}
  .article {{
    margin-bottom: 14px;
    page-break-inside: avoid;
  }}
  .article-title {{
    font-weight: bold;
    font-size: 11pt;
    margin-bottom: 6px;
    text-transform: uppercase;
  }}
  p {{
    margin: 0 0 6px 0;
    text-align: justify;
  }}
  .signing {{
    margin-top: 30px;
    display: flex;
    justify-content: space-between;
  }}
  .sign-block {{
    width: 45%;
  }}
  .sign-line {{
    border-top: 1px solid #000;
    margin-top: 40px;
    padding-top: 4px;
    font-size: 10pt;
  }}
  .po-meta {{
    font-size: 10pt;
    color: #333;
    margin-bottom: 14px;
  }}
</style>
</head>
<body>
  <div class="header-box">
    <div class="company-name">M/S. Passavant Energy &amp; Environment India Pvt. Ltd.</div>
    <div class="company-address">Navi Mumbai, India</div>
  </div>

  <div class="po-meta">PO Number: {po_number} | Site: {site_name}</div>

  <div class="attn">Kind Attn: {vendor_contact}</div>

  <div class="subject">Subject: LOI for {description}</div>

  <div class="intro">
    The contractor/buyer, Passavant Energy &amp; Environment India Pvt. Ltd. (PEEIPL),
    desiring that certain works pertaining to the project as described herein,
    has entered into discussions with M/s <strong>{vendor_name}</strong> (the Seller)
    and is pleased to issue this Letter of Intent subject to the terms and conditions set
    forth below.
  </div>

  {articles_html}

  <div style="margin-top:30px;">
    <p>Please sign and return a copy of this Letter of Intent as your acceptance of the
    above terms and conditions.</p>
  </div>

  <table style="width:100%;margin-top:40px;">
    <tr>
      <td style="width:50%;vertical-align:top;">
        <p><strong>For Passavant Energy &amp; Environment India Pvt. Ltd.</strong></p>
        <br><br><br>
        <div style="border-top:1px solid #000;padding-top:4px;font-size:10pt;">Authorised Signatory</div>
      </td>
      <td style="width:50%;vertical-align:top;padding-left:30px;">
        <p><strong>For {vendor_name}</strong></p>
        <br><br><br>
        <div style="border-top:1px solid #000;padding-top:4px;font-size:10pt;">Authorised Signatory</div>
      </td>
    </tr>
  </table>

</body>
</html>"""

    # ── Material PO document generators ───────────────────────────────────────
    # These produce the full PO document (PO header + Annexure-I line items +
    # Annexure-II Special Conditions + Annexure-III General Conditions) that
    # matches the exact layout from the physical PO form.

    @staticmethod
    def generate_material_po_pdf(po_data: dict[str, Any], line_items: list[dict]) -> bytes:
        """
        Generate the full Material Purchase Order PDF:
        Page 1  – PO header table (Order No, Vendor, amounts, T&C footer)
        Page 2  – Distribution / Prepared-by footer
        Page 3  – Annexure-I  (line items table)
        Page 4+ – Annexure-II (Special Conditions, clauses 1-7)
        Page N+ – Annexure-III (General Conditions, clauses 1-34)
        """
        try:
            from fpdf import FPDF
        except ImportError:
            raise RuntimeError("fpdf2 is not installed.")

        variables = LOIService._build_variables(po_data)
        total = float(po_data.get("total_amount", 0))
        total_words = _amount_to_words(total)
        vendor_name = po_data.get("vendor_name", "[VENDOR NAME]")
        po_number   = po_data.get("po_number", "DRAFT")
        po_date     = po_data.get("required_by", "")
        delivery_addr = po_data.get("delivery_address", "")

        class MatPDF(FPDF):
            def header(self):
                # Right-hand logo area
                self.set_font("Helvetica", "B", 13)
                self.set_xy(120, 8)
                self.cell(0, 6, "PASSAVANT", ln=True, align="R")
                self.set_font("Helvetica", "", 9)
                self.set_xy(120, 15)
                self.cell(0, 4, "Energy & Environment", ln=True, align="R")
                # Left header label
                self.set_font("Helvetica", "", 9)
                self.set_xy(10, 12)
                self.cell(60, 5, "Purchase Order", align="C")
                # QHSE / Date row
                self.set_font("Helvetica", "", 8)
                self.set_xy(120, 21)
                self.cell(40, 4, "QHSE Ref. No.")
                self.set_xy(160, 21)
                self.cell(0, 4, "", border="B")
                self.set_xy(120, 26)
                self.cell(40, 4, "Date:")
                self.set_xy(160, 26)
                self.cell(0, 4, po_date)
                self.ln(10)

            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 8)
                self.cell(0, 10, f"Page {self.page_no()} of {{nb}}", align="C")

        pdf = MatPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_margins(15, 35, 15)

        # ── PAGE 1: PO Header ────────────────────────────────────────────────
        pdf.add_page()

        def cell_label(w, h, txt, border="LRB"):
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(w, h, txt, border=border)

        def cell_val(w, h, txt, border="LRB", bold=False):
            pdf.set_font("Helvetica", "B" if bold else "", 8)
            pdf.cell(w, h, txt, border=border)

        # Company banner
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 8, "PASSAVANT ENERGY & ENVIRONMENT INDIA PRIVATE LIMITED",
                 border=1, ln=True, align="C")

        # Order No / Date row
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(30, 6, "Order No.:", border="LRT")
        pdf.cell(75, 6, po_number, border="RT")
        pdf.cell(0, 6, "", border="LRT", ln=True)
        pdf.cell(30, 6, "Date:", border="LRB")
        pdf.cell(75, 6, po_date, border="RB")
        pdf.cell(0, 6, "", border="LRB", ln=True)

        # Billing + Order To
        pdf.cell(105, 5, "BILLING AND CONSIGNEE ADDRESS:", border="LRT")
        pdf.cell(0, 5, "Order To:", border="LRT", ln=True)
        pdf.cell(105, 10, delivery_addr[:60], border="LRB")
        pdf.cell(0, 10, vendor_name[:60], border="LRB", ln=True)

        # Delivery
        pdf.cell(105, 5, "DELIVERY AND SHIP TO ADDRESS -", border="LRT")
        pdf.cell(0, 5, "Offer:", border="LRT", ln=True)
        pdf.cell(105, 10, delivery_addr[:60], border="LRB")
        pdf.cell(0, 10, "", border="LRB", ln=True)

        # Tel / Fax / Attn
        pdf.cell(35, 5, "Tel:", border="LRTB")
        pdf.cell(35, 5, "FAX:", border="RTB")
        pdf.cell(35, 5, "", border="RTB")
        pdf.cell(0, 5, "Attn:", border="LRTB", ln=True)

        # Items header
        cws = [14, 18, 68, 14, 14, 26, 26]
        hdrs = ["JOB\nCAT", "ITEM\nCODE", "DESCRIPTION", "UNIT", "QTY",
                "RATE\n(INR)", "AMOUNT\n(INR)"]
        pdf.set_font("Helvetica", "B", 7)
        for i, (w, h) in enumerate(zip(cws, hdrs)):
            pdf.cell(w, 8, h, border=1, align="C")
        pdf.ln()

        # Summary line item row
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(cws[0], 6, "", border="LR")
        pdf.cell(cws[1], 6, "", border="R")
        desc_short = po_data.get("description", "")[:40]
        pdf.cell(cws[2], 6, desc_short, border="R")
        pdf.cell(cws[3], 6, "LOT", border="R", align="C")
        pdf.cell(cws[4], 6, "1", border="R", align="C")
        pdf.cell(cws[5], 6, f"{total:,.0f}", border="R", align="R")
        pdf.cell(cws[6], 6, f"{total:,.0f}", border="R", align="R")
        pdf.ln()

        # T&C note inside item table
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(cws[0], 4, "", border="L")
        pdf.cell(cws[1], 4, "", border="")
        pdf.multi_cell(cws[2], 4,
            "TERMS & CONDITIONS:\nALL TERMS AND CONDITIONS SHALL BE APPLICABLE AS\nPER ANNEXURE-II & III.",
            border="")
        y_after = pdf.get_y()

        # Blank filler rows
        for _ in range(4):
            for w in cws:
                pdf.cell(w, 5, "", border="LR")
            pdf.ln()
        pdf.cell(0, 0, "", border="T", ln=True)

        # Totals / signatories
        half = sum(cws) / 2
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(half, 6, "PROJECT MANAGER", border="LRT")
        pdf.cell(half - cws[-1] - cws[-2], 6, "FINANCE MANAGER", border="LRT")
        pdf.cell(cws[-2], 6, "Total:", border="LRT", align="R")
        pdf.cell(cws[-1], 6, f"{total:,.0f}", border="LRT", align="R")
        pdf.ln()

        pdf.cell(half, 6, "", border="LRB")
        pdf.cell(half - cws[-1] - cws[-2], 6, "", border="LRB")
        pdf.cell(cws[-2], 6, "Disc Amount", border="LRB", align="R")
        pdf.cell(cws[-1], 6, "", border="LRB", align="R")
        pdf.ln()

        pdf.cell(half, 6, "PROCUREMENT\nMANAGER", border="LRT")
        pdf.cell(half - cws[-1] - cws[-2], 6, "MANAGING DIRECTOR", border="LRT")
        pdf.cell(cws[-2], 6, "Net Amount", border="LRT", align="R")
        pdf.cell(cws[-1], 6, f"{total:,.0f}", border="LRT", align="R")
        pdf.ln()
        pdf.cell(half, 6, "", border="LRB")
        pdf.cell(half - cws[-1] - cws[-2], 6, "", border="LRB")
        pdf.cell(cws[-2], 6, "", border="LRB")
        pdf.cell(cws[-1], 6, "", border="LRB")
        pdf.ln()

        # Amount in words
        words_line = f"INDIAN RUPEE {total_words.upper()}"
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(0, 6, words_line, border=1, ln=True, align="C")

        # Payment/delivery footer
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(35, 5, "Payment Terms", border="LRT")
        pdf.cell(0, 5, "", border="LRT", ln=True)
        pdf.cell(35, 5, "Del. Details", border="LRB")
        pdf.cell(0, 5, "AS MENTIONED ABOVE", border="LRB", ln=True)
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(105, 10,
            "PLEASE SIGN AND RETURN ACKNOWLEDGEMENT ORDER COPY AND CONFIRM\n"
            "DELIVERY BY RETURN\nStandard Terms and Conditions apply",
            border=1)
        pdf.cell(0, 10,
            "COPY OF ULTIMATE MANUFACTURERS FACTORY\n"
            "ORDER ACKNOWLEDGEMENT REQUIRED WITHIN\n7 DAYS OF ORDER DATE",
            border=1, ln=True)

        # ── PAGE 2: Distribution / Prepared-by ──────────────────────────────
        pdf.add_page()
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(105, 10, "Distribution: Supplier, Office, Site, Store and File", border=1)
        pdf.cell(0, 10, "Prepared by:", border=1, ln=True)

        # ── PAGE 3: Annexure-I Line Items ────────────────────────────────────
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "Annexure-I", ln=True, align="C")
        pdf.ln(2)

        col_ws = [14, 72, 18, 18, 28, 30]
        col_hs = ["S. No.", "Item Description", "Unit", "Qty.", "Rate", "Amount"]
        pdf.set_font("Helvetica", "B", 8)
        for w, h in zip(col_ws, col_hs):
            pdf.cell(w, 7, h, border=1, align="C")
        pdf.ln()

        subtotal = 0.0
        for i, item in enumerate(line_items, 1):
            rate = float(item.get("rate", 0))
            qty  = float(item.get("qty", 1))
            amt  = rate * qty
            subtotal += amt
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(col_ws[0], 10, str(i), border=1, align="C")
            pdf.cell(col_ws[1], 10, str(item.get("description", ""))[:60], border=1)
            pdf.cell(col_ws[2], 10, str(item.get("unit", "NOS")).upper(), border=1, align="C")
            pdf.cell(col_ws[3], 10, str(int(qty)), border=1, align="C")
            pdf.cell(col_ws[4], 10, f"{rate:,.0f}", border=1, align="R")
            pdf.cell(col_ws[5], 10, f"{amt:,.0f}", border=1, align="R")
            pdf.ln()

        # Sub-total / Freight / GST / Total rows
        pdf.set_font("Helvetica", "B", 8)
        label_w = sum(col_ws[:5])
        pdf.cell(label_w, 6, "Sub-Total", border=1, align="R")
        pdf.cell(col_ws[5], 6, f"{subtotal:,.0f}", border=1, align="R")
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(label_w, 6, "Freight (Packing & forwarding) & Transit Insurance", border=1)
        pdf.cell(col_ws[5], 6, "Incl", border=1, align="R")
        pdf.ln()
        pdf.cell(label_w, 6, "GST", border=1)
        pdf.cell(col_ws[5], 6, "Extra", border=1, align="R")
        pdf.ln()
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(label_w, 6, "Total Cost to Site", border=1, align="R")
        pdf.cell(col_ws[5], 6, f"{subtotal:,.0f}", border=1, align="R")
        pdf.ln()

        # ── ANNEXURE-II: Special Conditions ─────────────────────────────────
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "Annexure - II", ln=True, align="C")
        pdf.ln(1)
        pdf.set_font("Helvetica", "BU", 9)
        pdf.cell(0, 6, "SPECIAL CONDITIONS: -", ln=True)
        pdf.ln(1)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(0, 5, "DEFINITIONS", ln=True)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(0, 5, "CLIENT:", ln=True)
        pdf.cell(0, 5, "BUYER:", ln=True)
        pdf.cell(0, 5, "SELLER/VENDOR:", ln=True)
        pdf.ln(3)

        for clause in MATERIAL_SPECIAL_CONDITIONS:
            body = _safe_fill(clause["body"], variables)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 6, f"{clause['number']}.    {clause['title']}", ln=True)
            pdf.set_font("Helvetica", "", 8)
            for line in body.split("\n"):
                if line.strip().startswith("\u2022"):
                    pdf.cell(8, 5, "")
                    pdf.multi_cell(0, 5, line.strip())
                else:
                    pdf.multi_cell(0, 5, line)
            pdf.ln(2)

        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(0, 6, "For Passavant Energy & Environment India Private Limited.", ln=True)
        pdf.ln(10)
        pdf.cell(0, 5, "(Authorized Signatory)", ln=True)

        # ── ANNEXURE-III: General Conditions ────────────────────────────────
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "ANNEXURE-III", ln=True, align="C")
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, "GENERAL CONDITIONS", ln=True, align="C")
        pdf.ln(2)

        for clause in MATERIAL_GENERAL_CONDITIONS:
            body = _safe_fill(clause["body"], variables)
            pdf.set_font("Helvetica", "B", 9)
            heading = f"{clause['number']}.    {clause['title']}"
            pdf.multi_cell(0, 6, heading)
            pdf.set_font("Helvetica", "", 8)
            for line in body.split("\n"):
                if line.strip().startswith("\u2022"):
                    pdf.cell(8, 5, "")
                    pdf.multi_cell(0, 5, line.strip())
                else:
                    pdf.multi_cell(0, 5, line)
            pdf.ln(3)

        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(0, 6, "For Passavant Energy & Environment India Private Limited.", ln=True)
        pdf.ln(10)
        pdf.cell(0, 5, "(Authorized Signatory)", ln=True)

        buf = BytesIO()
        pdf.output(buf)
        return buf.getvalue()

    @staticmethod
    def generate_material_po_docx(po_data: dict[str, Any], line_items: list[dict]) -> bytes:
        """Generate the full Material PO as a Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError("python-docx is not installed.")

        variables = LOIService._build_variables(po_data)
        total = float(po_data.get("total_amount", 0))
        total_words = _amount_to_words(total)
        vendor_name = po_data.get("vendor_name", "[VENDOR NAME]")
        po_number   = po_data.get("po_number", "DRAFT")
        po_date     = po_data.get("required_by", "")
        delivery_addr = po_data.get("delivery_address", "")

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2)
            sec.right_margin  = Cm(2)

        def add_heading(text, level=1, bold=True, center=False):
            p = doc.add_paragraph()
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11 if level == 1 else 10)
            return p

        def add_body(text, bold=False):
            for line in text.split("\n"):
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.bold = bold
                r.font.size = Pt(9)
            return

        # ── PO Header ────────────────────────────────────────────────────────
        add_heading("PASSAVANT ENERGY & ENVIRONMENT INDIA PRIVATE LIMITED", center=True)
        doc.add_paragraph(f"Purchase Order No.: {po_number}    Date: {po_date}")
        doc.add_paragraph(f"Vendor: {vendor_name}")
        doc.add_paragraph(f"Delivery Address: {delivery_addr}")

        # Amount summary table
        t = doc.add_table(rows=3, cols=2)
        t.style = "Table Grid"
        t.cell(0, 0).text = "Description"
        t.cell(0, 1).text = po_data.get("description", "")[:80]
        t.cell(1, 0).text = "Total Amount (INR)"
        t.cell(1, 1).text = f"{total:,.0f}"
        t.cell(2, 0).text = "Amount in Words"
        t.cell(2, 1).text = f"INDIAN RUPEE {total_words.upper()}"
        doc.add_paragraph()

        # Terms note
        p = doc.add_paragraph()
        p.add_run(
            "TERMS & CONDITIONS: ALL TERMS AND CONDITIONS SHALL BE APPLICABLE "
            "AS PER ANNEXURE-II & III."
        ).bold = True
        doc.add_paragraph()

        # ── Annexure-I ───────────────────────────────────────────────────────
        add_heading("Annexure-I", center=True)
        doc.add_paragraph()
        n_cols = 6
        tbl = doc.add_table(rows=1, cols=n_cols)
        tbl.style = "Table Grid"
        hdrs = ["S. No.", "Item Description", "Unit", "Qty.", "Rate (INR)", "Amount (INR)"]
        for i, h in enumerate(hdrs):
            tbl.rows[0].cells[i].text = h
            tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        subtotal = 0.0
        for idx, item in enumerate(line_items, 1):
            rate = float(item.get("rate", 0))
            qty  = float(item.get("qty", 1))
            amt  = rate * qty
            subtotal += amt
            row = tbl.add_row().cells
            row[0].text = str(idx)
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("unit", "NOS")).upper()
            row[3].text = str(int(qty))
            row[4].text = f"{rate:,.0f}"
            row[5].text = f"{amt:,.0f}"

        for label, val in [
            ("Sub-Total", f"{subtotal:,.0f}"),
            ("Freight (Packing & forwarding) & Transit Insurance", "Incl"),
            ("GST", "Extra"),
            ("Total Cost to Site", f"{subtotal:,.0f}"),
        ]:
            row = tbl.add_row().cells
            row[1].merge(row[4])
            row[1].text = label
            row[5].text = val
        doc.add_paragraph()

        # ── Annexure-II: Special Conditions ─────────────────────────────────
        add_heading("Annexure - II", center=True)
        add_heading("SPECIAL CONDITIONS: -", bold=True)
        add_heading("DEFINITIONS", bold=True)
        doc.add_paragraph("CLIENT:\nBUYER:\nSELLER/VENDOR:")
        doc.add_paragraph()

        for clause in MATERIAL_SPECIAL_CONDITIONS:
            body = _safe_fill(clause["body"], variables)
            p = doc.add_paragraph()
            p.add_run(f"{clause['number']}.    {clause['title']}").bold = True
            add_body(body)
            doc.add_paragraph()

        doc.add_paragraph("For Passavant Energy & Environment India Private Limited.")
        doc.add_paragraph("\n\n(Authorized Signatory)")
        doc.add_page_break()

        # ── Annexure-III: General Conditions ────────────────────────────────
        add_heading("ANNEXURE-III", center=True)
        add_heading("GENERAL CONDITIONS", center=True)
        doc.add_paragraph()

        for clause in MATERIAL_GENERAL_CONDITIONS:
            body = _safe_fill(clause["body"], variables)
            p = doc.add_paragraph()
            p.add_run(f"{clause['number']}.    {clause['title']}").bold = True
            add_body(body)
            doc.add_paragraph()

        doc.add_paragraph("For Passavant Energy & Environment India Private Limited.")
        doc.add_paragraph("\n\n(Authorized Signatory)")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
